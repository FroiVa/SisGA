# export_utils.py
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime, date
from django.http import HttpResponse
import io
from django.contrib.auth.models import User
from .models import Area, Estado


def exportar_incidencias_docx(tabla_datos, dias, area_responsable, fecha_inicio, fecha_fin, mes_actual,
                              areas_hijas=None):
    """
    Exporta las incidencias a un documento DOCX con formato específico
    """
    # Crear documento
    doc = Document()

    # ==================== CONFIGURAR ORIENTACIÓN HORIZONTAL ====================
    # Cambiar orientación a horizontal (paisaje) para toda la sección
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)  # Ancho de página A4 en horizontal
    section.page_height = Cm(21.0)  # Alto de página A4 en horizontal

    # Configurar márgenes (opcional, para mejor ajuste)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    # ===========================================================================

    # _________________________________________________________________________
    # Añadir título y encabezado
    # _________________________________________________________________________

    # Período
    p = doc.add_paragraph()
    p.add_run(f"Período: {dias[0].day}/01/2026 al {dias[-1].day}/01/2026").bold = True
    p.add_run(" Empresa: Universidad de La Habana")

    # Generado
    p = doc.add_paragraph()
    generated_date = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    p.add_run(f"Generado: {generated_date}").bold = True
    p.add_run(f"Unidad: {area_responsable.area.nombre}")

    # Organismo
    p = doc.add_paragraph()
    p.add_run("Organismo: MES").bold = True
    p.add_run(f"Confeccionado por: {area_responsable.usuario.first_name} {area_responsable.usuario.last_name}")

    # Aprobado por
    doc.add_paragraph("Aprobado por:")
    doc.add_paragraph("_" * 55)

    # Espacio
    doc.add_paragraph()
    # ________________________________________________________________________________
    # Crear tabla principal 5x17 (filas x columnas)
    # ________________________________________________________________________________
    mitad = len(dias)//2
    parte1 = dias[:mitad]  # [1,..., 15]
    parte2 = dias[mitad:]  # [16,...,31]
    table = doc.add_table(rows=2, cols=len(parte2)+2)
    table.style = 'Table Grid'



    hdr_cells_row1 = table.rows[0].cells
    hdr_cells_row1[0].width = Cm(3)
    hdr_cells_row1[0].text = 'Expte'
    hdr_cells_row1[1].width = Cm(5)
    hdr_cells_row1[1].text = 'Nombre y Apellidos'
    for i in range(2, len(parte1) + 2):
        hdr_cells_row1[i].width = Cm(1.5)
        hdr_cells_row1[i].text = f"{parte1[i-2].day}"
        print(parte1[i-2].day)

    hdr_cells_row2 = table.rows[1].cells
    hdr_cells_row2[0].text = ''
    hdr_cells_row2[1].text = ''
    for i in range(2, len(parte1) + 3):
        hdr_cells_row2[i].text = f"{parte2[i-2].day}"
        print(parte2[i-2].day)

    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))


    for filas in tabla_datos:
        mitad_estado = len(filas['dias']) // 2
        mitad1 =filas['dias'][:mitad_estado]
        mitad2 =filas['dias'][mitad_estado:]
        # Crear dos filas
        row_cells = table.add_row().cells
        row_cells1 = table.add_row().cells

        # Mostrando primera mitad del período
        row_cells[0].width = Cm(3)
        row_cells[0].text = filas['expte']
        row_cells[1].width = Cm(5)
        row_cells[1].text = filas['empleado']
        for i, dia_data in enumerate(mitad1, start=2):
            row_cells[i].width = Cm(1.5)
            row_cells[i].text = dia_data['estado'].clave_id

        # Mostrando segunda mitad del período
        row_cells1[0].width = Cm(3)
        row_cells1[0].text = ''
        row_cells1[1].width = Cm(5)
        row_cells1[1].text = ''
        for i, dia_data in enumerate(mitad2, start=2):
            row_cells1[i].width = Cm(1.5)
            row_cells1[i].text = dia_data['estado'].clave_id
        row_cells[0].merge(row_cells1[0])
        row_cells[1].merge(row_cells1[1])


    return doc

# def aplicar_estilos_documento(doc):
#     """Aplica estilos personalizados al documento"""
#     # Configurar estilos de párrafo normales
#     style = doc.styles['Normal']
#     style.font.name = 'Arial'
#     style.font.size = Pt(11)
#
#     # Estilo para encabezados
#     for i in range(1, 4):
#         style = doc.styles[f'Heading {i}']
#         style.font.name = 'Arial'
#         style.font.size = Pt(14 - i * 2)
#         style.font.bold = True
#
#     return doc
#
#
# def generar_respuesta_docx(doc, nombre_archivo):
#     """
#     Genera una respuesta HTTP con el documento DOCX
#     """
#     # Aplicar estilos
#     doc = aplicar_estilos_documento(doc)
#
#     # Guardar en buffer
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#
#     # Crear respuesta
#     response = HttpResponse(
#         buffer.getvalue(),
#         content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#     )
#     response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}.docx"'
#
#     return response
